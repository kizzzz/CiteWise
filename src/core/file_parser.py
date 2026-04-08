"""统一文件解析入口 — 支持 PDF/DOCX/MD/TXT/XLSX"""
import os
import uuid
import logging
from typing import Optional

from config.settings import PAPERS_DIR

logger = logging.getLogger(__name__)

# Supported extensions
SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".md", ".txt", ".xlsx", ".xls"}


def get_file_extension(filename: str) -> str:
    """获取文件后缀（小写）"""
    return os.path.splitext(filename)[1].lower()


def is_supported(filename: str) -> bool:
    return get_file_extension(filename) in SUPPORTED_EXTENSIONS


def parse_file(filepath: str, filename: Optional[str] = None) -> dict:
    """统一文件解析入口

    Returns: dict with keys: paper_id, title, authors, year, sections, raw_text, figures
    """
    fname = filename or os.path.basename(filepath)
    ext = get_file_extension(fname)

    if ext == ".pdf":
        return _parse_pdf(filepath, fname)
    elif ext in (".doc", ".docx"):
        return _parse_docx(filepath, fname)
    elif ext in (".md", ".txt"):
        return _parse_text(filepath, fname)
    elif ext in (".xlsx", ".xls"):
        return _parse_xlsx(filepath, fname)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


def _parse_pdf(filepath: str, filename: str) -> dict:
    """PDF 解析 — 复用 rag.parse_pdf"""
    from src.core.rag import parse_pdf
    return parse_pdf(filepath)


def _parse_docx(filepath: str, filename: str) -> dict:
    """DOCX 解析 — python-docx"""
    paper_id = f"paper_{uuid.uuid4().hex[:8]}"
    title = os.path.splitext(filename)[0]
    authors = ""
    year = 0

    try:
        from docx import Document
        doc = Document(filepath)

        # Try to extract metadata
        core_props = doc.core_properties
        if core_props.title:
            title = core_props.title
        if core_props.author:
            authors = core_props.author
        if core_props.created:
            year = core_props.created.year

        # Extract paragraphs
        full_text_parts = []
        sections = []
        current_section = {"title": "全文", "text": "", "tables": []}

        import re
        heading_pattern = re.compile(r'^(\d+(?:\.\d+)*)\s+(.+)|^([一二三四五六七八九十]+[、．.])\s*(.+)', re.UNICODE)

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if heading
            if para.style and para.style.name and "Heading" in para.style.name:
                if current_section["text"].strip():
                    sections.append(current_section.copy())
                current_section = {"title": text, "text": "", "tables": []}
            elif heading_pattern.match(text) and len(text) < 80:
                if current_section["text"].strip():
                    sections.append(current_section.copy())
                current_section = {"title": text, "text": "", "tables": []}
            else:
                current_section["text"] += text + "\n"
                full_text_parts.append(text)

        # Also extract tables
        for table in doc.tables:
            table_md = _docx_table_to_markdown(table)
            if table_md:
                current_section["tables"].append({"content": table_md})

        if current_section["text"].strip():
            sections.append(current_section)

        raw_text = "\n".join(full_text_parts)

        # If no sections detected, use full text as one section
        if not sections:
            sections = [{"title": "全文", "text": raw_text, "tables": []}]

    except Exception as e:
        logger.error(f"DOCX 解析失败: {e}")
        return {
            "paper_id": paper_id,
            "filename": filename,
            "title": title,
            "authors": authors,
            "year": year,
            "sections": [],
            "raw_text": "",
            "figures": [],
            "error": str(e),
        }

    return {
        "paper_id": paper_id,
        "filename": filename,
        "title": title,
        "authors": authors,
        "year": year,
        "sections": sections,
        "raw_text": raw_text,
        "figures": [],
    }


def _parse_text(filepath: str, filename: str) -> dict:
    """TXT/MD 解析 — 直接读取文本"""
    paper_id = f"paper_{uuid.uuid4().hex[:8]}"
    title = os.path.splitext(filename)[0]

    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            raw_text = f.read()

        # For markdown, try to split by headings
        import re
        sections = []
        current_section = {"title": "全文", "text": "", "tables": []}

        md_heading = re.compile(r'^(#{1,6})\s+(.+)', re.MULTILINE)
        lines = raw_text.split('\n')
        buffer = []

        for line in lines:
            match = md_heading.match(line)
            if match:
                if buffer:
                    current_section["text"] += "\n".join(buffer) + "\n"
                    buffer = []
                if current_section["text"].strip():
                    sections.append(current_section.copy())
                current_section = {"title": match.group(2).strip(), "text": "", "tables": []}
            else:
                buffer.append(line)

        if buffer:
            current_section["text"] += "\n".join(buffer) + "\n"
        if current_section["text"].strip():
            sections.append(current_section)

        if not sections:
            sections = [{"title": "全文", "text": raw_text, "tables": []}]

    except Exception as e:
        logger.error(f"文本文件解析失败: {e}")
        return {
            "paper_id": paper_id,
            "filename": filename,
            "title": title,
            "authors": "",
            "year": 0,
            "sections": [],
            "raw_text": "",
            "figures": [],
            "error": str(e),
        }

    return {
        "paper_id": paper_id,
        "filename": filename,
        "title": title,
        "authors": "",
        "year": 0,
        "sections": sections,
        "raw_text": raw_text,
        "figures": [],
    }


def _parse_xlsx(filepath: str, filename: str) -> dict:
    """XLSX 解析 — openpyxl 按行提取"""
    paper_id = f"paper_{uuid.uuid4().hex[:8]}"
    title = os.path.splitext(filename)[0]

    try:
        from openpyxl import load_workbook
        wb = load_workbook(filepath, read_only=True, data_only=True)

        sections = []
        full_text_parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):  # skip fully empty rows
                    rows.append(cells)

            if not rows:
                continue

            # Convert to text: first row as header, rest as data
            text_parts = [f"## Sheet: {sheet_name}"]
            if rows:
                header = rows[0]
                text_parts.append("| " + " | ".join(header) + " |")
                text_parts.append("| " + " | ".join("---" for _ in header) + " |")
                for row in rows[1:]:
                    # Pad row to match header length
                    padded = row + [""] * (len(header) - len(row))
                    text_parts.append("| " + " | ".join(padded[:len(header)]) + " |")

            section_text = "\n".join(text_parts)
            sections.append({"title": sheet_name, "text": section_text, "tables": []})
            full_text_parts.append(section_text)

        wb.close()

        raw_text = "\n\n".join(full_text_parts)
        if not sections:
            sections = [{"title": "全文", "text": raw_text, "tables": []}]

    except Exception as e:
        logger.error(f"XLSX 解析失败: {e}")
        return {
            "paper_id": paper_id,
            "filename": filename,
            "title": title,
            "authors": "",
            "year": 0,
            "sections": [],
            "raw_text": "",
            "figures": [],
            "error": str(e),
        }

    return {
        "paper_id": paper_id,
        "filename": filename,
        "title": title,
        "authors": "",
        "year": 0,
        "sections": sections,
        "raw_text": raw_text,
        "figures": [],
    }


def _docx_table_to_markdown(table) -> str:
    """Convert python-docx table to markdown"""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    header = "| " + " | ".join(rows[0]) + " |"
    separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
    data_rows = ["| " + " | ".join(r) + " |" for r in rows[1:]]
    return header + "\n" + separator + "\n" + "\n".join(data_rows)
